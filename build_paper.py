#!/usr/bin/env python3
"""
Build RobustCrop paper as DOCX via pandoc.
Reads paper_draft_v2.md, injects proper figures/math/tables, outputs .docx.
"""

import subprocess, re, os, sys

FIGURES_DIR = "results/figures"
OUTPUT_DOCX = "RobustCrop_Paper_v2.docx"
OUTPUT_PDF  = "RobustCrop_Paper_v2.pdf"

# ── Figure mapping (appendix figure numbers → actual files) ──────────────
FIG_MAP = {
    1:  "01_primary_distributions.png",
    2:  "02_secondary_distributions.png",
    3:  "03_primary_correlation.png",
    4:  "04_secondary_correlation.png",
    5:  "05_class_distributions.png",
    6:  "06_degradation_comparison.png",
    7:  "07_shared_features_comparison.png",
    8:  "08_feature_selection.png",
    9:  "09_feature_selection_secondary.png",
    10: "10_cv_comparison.png",
    11: "11_shap_RandomForest.png",
    12: "12_robustness.png",
    13: "13_calibration.png",
    14: "14_per_class_heatmap.png",
}

def build_pandoc_markdown(src_path: str) -> str:
    """Transform paper_draft_v2.md into pandoc-flavoured markdown with:
    - YAML front matter (for docx metadata)
    - Proper figure image blocks with captions
    - Proper display math blocks
    - Proper table formatting
    """
    with open(src_path) as f:
        text = f.read()

    # ── 1. YAML front matter ────────────────────────────────────────────
    yaml_header = """---
title: "RobustCrop: A Leak-Free Machine Learning Pipeline for Crop Recommendation with Sensor Degradation Analysis and Cross-Dataset Feature Consistency"
author:
  - "Anuradha Brijwal"
  - "Praveena Chaturvedi"
date: "2026"
institute:
  - "Research Scholar, Department of Computer Science, Kanya Gurukul Campus Dehradun"
  - "Professor, Department of Computer Science, Kanya Gurukul Campus Dehradun"
subtitle: "Gurukul Kangri (Deemed to be University), Haridwar, Uttarakhand, India"
geometry: "margin=1in"
fontsize: 11pt
reference-section-title: "References"
# csl: "apa.csl"
link-citations: true
---

"""

    # ── 2. Remove the old title block (everything before ## Abstract) ───
    abstract_idx = text.find("## Abstract")
    if abstract_idx == -1:
        print("ERROR: Could not find '## Abstract' in source")
        sys.exit(1)
    text = text[abstract_idx:]

    # ── 3. Replace display math:  $$...$$  stays as-is for pandoc ──────
    # Pandoc natively handles $$...$$ for display math and $...$ for inline.
    # But we need to make sure our LaTeX is clean.
    # Fix common issues:
    #   \mathbf, \mathbb, \boldsymbol → keep
    #   \text{} → \mathrm{} (pandoc handles \text in math mode)
    #   \sum, \frac, \left, \right → keep

    # ── 4. Inject figure images after the "Appendix B: Figures" table ───
    # We'll add proper pandoc figure blocks: ![Caption](path){#fig:N}
    fig_block = "\n\n"
    for num, fname in sorted(FIG_MAP.items()):
        path = os.path.join(FIGURES_DIR, fname)
        if os.path.exists(path):
            # Build caption from the appendix table
            captions = {
                1: "Primary dataset feature distributions (N, P, K, Temperature, Humidity, pH, Rainfall).",
                2: "Secondary dataset feature distributions (12 soil fertility features).",
                3: "Primary dataset correlation heatmap showing inter-feature relationships.",
                4: "Secondary dataset correlation heatmap for soil fertility features.",
                5: "Class distribution comparison: 22 balanced crop classes (primary) vs. 3 imbalanced fertility levels (secondary).",
                6: "Sensor degradation effect on feature distributions across deployment scenarios.",
                7: "Shared feature space comparison (N, P, K, pH) across datasets.",
                8: "Feature selection methods — Primary dataset: six methods plus consensus ranking.",
                9: "Feature selection methods — Secondary dataset: six methods plus consensus ranking.",
                10: "Cross-validation accuracy comparison across all 10 classifiers on the primary dataset.",
                11: "SHAP feature importance for the proposed Random Forest classifier.",
                12: "Model robustness under sensor degradation: accuracy vs. deployment days without recalibration.",
                13: "Calibration curves for top classifiers showing reliability of predicted probabilities.",
                14: "Per-class F1 score heatmap across 22 crop classes and 10 classifiers.",
            }
            cap = captions.get(num, f"Figure {num}")
            fig_block += f"![{cap}]({path}){{#fig:{num} width=100%}}\n\n"

    # Insert figures right after the Appendix B figure table
    appendix_b_marker = "## Appendix B: Figures"
    if appendix_b_marker in text:
        # Find the end of the appendix B table
        idx = text.find(appendix_b_marker)
        # Find next ## or end
        next_section = text.find("\n## ", idx + len(appendix_b_marker))
        if next_section == -1:
            next_section = len(text)
        # Insert figures before the next section
        text = text[:next_section] + fig_block + text[next_section:]

    # ── 5. Also inject key figures inline where they are referenced ─────
    # After "SHAP analysis" section (5.7), inject SHAP figure
    # After "Sensor Degradation" section (5.6), inject robustness figure
    # After "Calibration" section (6.4), inject calibration figure
    inline_figures = {
        "### 5.1 Real-World Secondary Dataset Results": [
            (f"![Cross-validation accuracy comparison across classifiers on the primary dataset.]({FIGURES_DIR}/10_cv_comparison.png){{width=100%}}\n\n", "after"),
        ],
        "### 5.3 Feature Selection Analysis": [
            (f"![Feature selection consensus ranking for the primary dataset.]({FIGURES_DIR}/08_feature_selection.png){{width=100%}}\n\n", "after"),
            (f"![Feature selection consensus ranking for the secondary dataset.]({FIGURES_DIR}/09_feature_selection_secondary.png){{width=100%}}\n\n", "after"),
        ],
        "### 5.6 Sensor Degradation Robustness": [
            (f"![Model robustness under sensor degradation: accuracy vs. deployment days.]({FIGURES_DIR}/12_robustness.png){{width=100%}}\n\n", "after"),
        ],
        "### 5.7 SHAP Explainability Analysis": [
            (f"![SHAP feature importance for the proposed Random Forest classifier.]({FIGURES_DIR}/11_shap_RandomForest.png){{width=100%}}\n\n", "after"),
        ],
    }

    for marker, inserts in inline_figures.items():
        if marker in text:
            idx = text.find(marker)
            # Find end of this section (next ### or ##)
            end_marker = text.find("\n### ", idx + len(marker))
            if end_marker == -1:
                end_marker = text.find("\n## ", idx + len(marker))
            if end_marker == -1:
                end_marker = len(text)
            for fig_md, position in inserts:
                if position == "after":
                    # Insert before the next section
                    text = text[:end_marker] + "\n" + fig_md + text[end_marker:]
                    # Recalculate end_marker for next insert
                    end_marker = text.find("\n### ", end_marker + len(fig_md))
                    if end_marker == -1:
                        end_marker = text.find("\n## ", end_marker + len(fig_md))
                    if end_marker == -1:
                        end_marker = len(text)

    # ── 6. Fix inline math: ensure proper $...$ delimiters ──────────────
    # The paper uses Unicode math symbols (χ², κ, etc.) which pandoc handles fine.
    # But we should also convert key formulas to proper LaTeX inline math.
    # Replace specific Unicode math with LaTeX equivalents in display math contexts
    math_replacements = [
        # The cross-dataset consistency formula
        (
            r"$$\text{Consistency}_j = 1 - |s_j^{\text{primary}} - s_j^{\text{secondary}}|$$",
            r"$$\mathrm{Consistency}_j = 1 - \left| s_j^{\mathrm{primary}} - s_j^{\mathrm{secondary}} \right|$$"
        ),
        # Class weighting formula
        (
            r"$$w_c = \frac{N}{C \cdot n_c}$$",
            r"$$w_c = \frac{N}{C \cdot n_c}$$"
        ),
        # Scaling formula
        (
            r"$$\tilde{\mathbf{x}}_i = \frac{\mathbf{x}_i - \boldsymbol{\mu}_k}{\boldsymbol{\sigma}_k}$$",
            r"$$\tilde{\mathbf{x}}_i = \frac{\mathbf{x}_i - \boldsymbol{\mu}_k}{\boldsymbol{\sigma}_k}$$"
        ),
        # MI formula
        (
            r"$$I(\mathbf{x}_j; y) = H(y) - H(y | \mathbf{x}_j)$$",
            r"$$I(\mathbf{x}_j; y) = H(y) - H(y \mid \mathbf{x}_j)$$"
        ),
        # Friedman statistic
        (
            r"$$\chi^2_F = \frac{12N}{k(k+1)}\left[\sum_j R_j^2 - \frac{k(k+1)^2}{4}\right]$$",
            r"$$\chi^2_F = \frac{12N}{k(k+1)}\left[\sum_j R_j^2 - \frac{k(k+1)^2}{4}\right]$$"
        ),
        # Pipeline math notation
        (
            r"$\mathcal{D} = \{(\mathbf{x}_i, y_i)\}_{i=1}^{N}$",
            r"$\mathcal{D} = \{(\mathbf{x}_i, y_i)\}_{i=1}^{N}$"
        ),
        (
            r"$\mathbf{x}_i \in \mathbb{R}^d$",
            r"$\mathbf{x}_i \in \mathbb{R}^d$"
        ),
        (
            r"$y_i \in \{1, \ldots, C\}$",
            r"$y_i \in \{1, \ldots, C\}$"
        ),
        (
            r"$\mathcal{D} = \bigcup_{k=1}^{K} \mathcal{F}_k$",
            r"$\mathcal{D} = \bigcup_{k=1}^{K} \mathcal{F}_k$"
        ),
        (
            r"$\mathcal{D} \setminus \mathcal{F}_k$",
            r"$\mathcal{D} \setminus \mathcal{F}_k$"
        ),
        (
            r"$\boldsymbol{\mu}_k$",
            r"$\boldsymbol{\mu}_k$"
        ),
        (
            r"$\boldsymbol{\sigma}_k$",
            r"$\boldsymbol{\sigma}_k$"
        ),
        (
            r"$I(\mathbf{x}_j; y)$",
            r"$I(\mathbf{x}_j; y)$"
        ),
        (
            r"$\bar{s}_j = \frac{1}{6}\sum_{m=1}^{6} s_j^{(m)}$",
            r"$\bar{s}_j = \frac{1}{6}\sum_{m=1}^{6} s_j^{(m)}$"
        ),
        (
            r"$R_j$",
            r"$R_j$"
        ),
    ]

    for old, new in math_replacements:
        text = text.replace(old, new)

    # ── 7. Fix Greek letters and special chars for pandoc ───────────────
    # Replace Unicode Greek with LaTeX math where needed in text
    # κ (kappa) → $\kappa$ when used as a metric
    # χ² → $\chi^2$ when used in Friedman test
    # These are already used inline in the text and pandoc handles Unicode fine
    # in docx output, so we leave them as-is.

    # ── 8. Clean up any stray formatting ────────────────────────────────
    # Remove any HTML comments
    text = re.sub(r'<!--.*?-->', '', text, flags=re.DOTALL)
    # Ensure blank lines before headers
    text = re.sub(r'\n(##)', r'\n\n\1', text)
    # Ensure blank lines before lists
    text = re.sub(r'\n(\d+\.)', r'\n\n\1', text)
    # Fix multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)

    return yaml_header + text


def write_pandoc_md(md_text: str, path: str = "paper_pandoc.md"):
    with open(path, "w") as f:
        f.write(md_text)
    print(f"  ✓ Pandoc markdown: {path} ({len(md_text)} chars)")
    return path


def build_docx(md_path: str, docx_path: str):
    """Build DOCX via pandoc with math as OMML (native Word equations)."""
    cmd = [
        "pandoc", md_path,
        "-o", docx_path,
        "--from", "markdown+tex_math_dollars+implicit_figures+table_captions+smart",
        "--to", "docx",
        "--mathml",                    # MathML → Word converts to native equations
        "--standalone",
        "--toc",                       # Table of contents
        "--toc-depth=3",
        "--number-sections",
        "--reference-doc=reference.docx" if os.path.exists("reference.docx") else None,
        "--resource-path=.:results/figures",
        "--highlight-style=tango",
    ]
    cmd = [c for c in cmd if c is not None]
    print(f"  Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  STDERR: {result.stderr}")
        # Don't fail, just warn - pandoc may produce output with warnings
    if os.path.exists(docx_path):
        size = os.path.getsize(docx_path)
        print(f"  ✓ DOCX: {docx_path} ({size/1024:.1f} KB)")
    else:
        print(f"  ✗ DOCX not created!")
        sys.exit(1)


def build_pdf(md_path: str, pdf_path: str):
    """Build PDF via pandoc + xelatex (if available)."""
    # Check for LaTeX
    try:
        subprocess.run(["xelatex", "--version"], capture_output=True, check=True)
    except (FileNotFoundError, subprocess.CalledProcessError):
        print("  ⚠ xelatex not available, skipping PDF generation")
        return

    cmd = [
        "pandoc", md_path,
        "-o", pdf_path,
        "--from", "markdown+tex_math_dollars+implicit_figures+table_captions+smart",
        "--pdf-engine=xelatex",
        "--standalone",
        "--toc",
        "--toc-depth=3",
        "--number-sections",
        "--resource-path=.:results/figures",
        "--highlight-style=tango",
        "-V", "geometry:margin=1in",
        "-V", "fontsize=11pt",
        "-V", "documentclass=article",
        "-V", "classoption=a4paper",
    ]
    print(f"  Running: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"  STDERR: {result.stderr[:500]}")
    if os.path.exists(pdf_path):
        size = os.path.getsize(pdf_path)
        print(f"  ✓ PDF: {pdf_path} ({size/1024:.1f} KB)")
    else:
        print(f"  ⚠ PDF not created (LaTeX may need additional packages)")


def verify_docx(docx_path: str):
    """Verify the generated DOCX has proper content."""
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ python-docx not available, skipping verification")
        return

    doc = Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = '\n'.join(paragraphs)

    checks = {
        "Title present":            "RobustCrop" in full_text,
        "Authors present":          "Anuradha" in full_text and "Chaturvedi" in full_text,
        "Abstract present":         "Machine learning for crop recommendation promises" in full_text,
        "Farmer intro":             "smallholder farmer" in full_text,
        "91.25% result":            "91.25%" in full_text,
        "99.50% result":            "99.50%" in full_text,
        "Friedman χ²":              "32.32" in full_text,
        "References present":       "Niculescu-Mizil" in full_text or "Zadrozny" in full_text,
        "No BalWeightWrapper in abstract": "BalWeightWrapper" not in full_text[:3000],
        "6 keywords":               "Explainable AI" in full_text,
    }

    # Check for images
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    image_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    checks[f"Images embedded ({image_count})"] = image_count >= 10

    # Check for math (OMML elements)
    import xml.etree.ElementTree as ET
    ns = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
    math_count = 0
    for p in doc.paragraphs:
        for run in p.runs:
            if run._element.findall('.//m:oMath', ns):
                math_count += 1
    checks[f"Math equations ({math_count})"] = math_count >= 3

    print("\n  ── DOCX Verification ──")
    all_ok = True
    for name, ok in checks.items():
        status = "✓" if ok else "✗"
        print(f"    {status} {name}")
        if not ok:
            all_ok = False
    return all_ok


if __name__ == "__main__":
    print("=" * 60)
    print("Building RobustCrop Paper v2")
    print("=" * 60)

    # Step 1: Build pandoc markdown
    print("\n[1/4] Building pandoc markdown...")
    md_text = build_pandoc_markdown("paper_draft_v2.md")
    md_path = write_pandoc_md(md_text)

    # Step 2: Generate DOCX
    print("\n[2/4] Generating DOCX...")
    build_docx(md_path, OUTPUT_DOCX)

    # Step 3: Generate PDF (optional)
    print("\n[3/4] Generating PDF (if LaTeX available)...")
    build_pdf(md_path, OUTPUT_PDF)

    # Step 4: Verify
    print("\n[4/4] Verifying DOCX...")
    ok = verify_docx(OUTPUT_DOCX)

    # Cleanup
    if os.path.exists(md_path) and md_path != "paper_draft_v2.md":
        os.remove(md_path)

    print("\n" + "=" * 60)
    if ok:
        print("✓ Build complete — all checks passed")
    else:
        print("⚠ Build complete — some checks failed, review output above")
    print("=" * 60)
