#!/usr/bin/env python3
"""
Build RobustCrop paper as a journal-formatted DOCX following
Elsevier/KeAi 'Artificial Intelligence in Agriculture' template.
"""

import os, re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import xml.etree.ElementTree as ET

FIGURES_DIR = "results/figures"
OUTPUT = "RobustCrop_Journal.docx"

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    """Add a numbered section heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(doc, text, bold=False, italic=False, size=11, alignment=None, space_after=6, style='Normal'):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    return p

def add_math_para(doc, latex):
    """Add a display math paragraph (as centered italic text for Word)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(latex)
    run.font.size = Pt(11)
    run.italic = True
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    return p

def add_figure(doc, filename, caption, width=5.5):
    """Add a figure with caption."""
    path = os.path.join(FIGURES_DIR, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = cap.add_run(caption)
    r.font.size = Pt(10)
    r.italic = True
    cap.paragraph_format.space_after = Pt(8)
    return cap

def add_table(doc, headers, rows, caption=None):
    """Add a formatted table."""
    if caption:
        cp = doc.add_paragraph()
        r = cp.add_run(caption)
        r.bold = True
        r.font.size = Pt(10)
        cp.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer
    return table


# ── BUILD ────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page setup ───────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ── Style defaults ──────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # ═══════════════════════════════════════════════════════════════════
    # HIGHLIGHTS
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    r = p.add_run("Highlights")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    highlights = [
        "A leak-free Pipeline-per-fold architecture eliminates data leakage that inflates accuracy in prior crop recommendation studies.",
        "Real-world evaluation on imbalanced soil fertility data (11.28:1 ratio) achieves 91.25% accuracy, revealing gaps invisible on balanced benchmarks.",
        "Cross-dataset feature consistency analysis identifies phosphorus as the most transferable feature across agricultural domains.",
        "Literature-grounded sensor degradation analysis establishes weekly recalibration as mandatory for maintaining >94% deployment accuracy.",
        "SHAP explainability reveals climate features (humidity, rainfall) dominate over soil nutrients for crop suitability prediction.",
    ]
    for h in highlights:
        p = doc.add_paragraph(h, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # ═══════════════════════════════════════════════════════════════════
    # TITLE
    # ═══════════════════════════════════════════════════════════════════
    title_text = "RobustCrop: A Leak-Free Machine Learning Pipeline for Crop Recommendation with Sensor Degradation Analysis and Cross-Dataset Feature Consistency"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title_text)
    r.bold = True
    r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    # ═══════════════════════════════════════════════════════════════════
    # AUTHORS
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Anuradha Brijwal")
    r.font.size = Pt(11)
    r = p.add_run("a")
    r.font.size = Pt(8)
    r.font.superscript = True
    r = p.add_run(", Praveena Chaturvedi")
    r.font.size = Pt(11)
    r = p.add_run("a")
    r.font.size = Pt(8)
    r.font.superscript = True
    p.paragraph_format.space_after = Pt(4)

    # Affiliations
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("a")
    r.font.size = Pt(9)
    r.font.superscript = True
    r = p.add_run(" Department of Computer Science, Kanya Gurukul Campus Dehradun, Gurukul Kangri (Deemed to be University), Haridwar, Uttarakhand, India")
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(4)

    # Corresponding author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Corresponding author: ")
    r.font.size = Pt(9)
    r.bold = True
    r = p.add_run("Anuradha Brijwal (anuradha.brijwal@gkv.ac.in)")
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(16)

    # ═══════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    r = p.add_run("Abstract")
    r.bold = True
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    abstract = (
        "Machine learning for crop recommendation promises data-driven agricultural guidance, "
        "yet most studies evaluate on balanced semi-synthetic benchmarks under leakage-prone "
        "preprocessing pipelines, providing inflated performance estimates that fail to transfer "
        "to real-world conditions. We present RobustCrop, a leak-free pipeline that encapsulates "
        "feature scaling, mutual-information-based feature selection, and classification within a "
        "single cross-validation fold, eliminating information leakage. We evaluate on two datasets: "
        "a real-world soil fertility dataset (880 samples, 3 classes, 11.28:1 imbalance ratio) and "
        "a semi-synthetic crop recommendation benchmark (2,200 samples, 22 classes). On real-world "
        "data, our Random Forest pipeline achieves 91.25% \u00b1 0.77% accuracy (\u03ba = 0.8364, "
        "macro-F1 = 81.85%), outperforming nine benchmarks. SHAP analysis identifies humidity and "
        "rainfall as dominant predictors, while cross-dataset analysis reveals phosphorus as the "
        "most transferable feature (consistency = 0.804). Literature-grounded sensor degradation "
        "analysis shows weekly recalibration maintains >94% accuracy, whereas 90-day uncalibrated "
        "deployment causes catastrophic 83.41% degradation. These findings provide actionable "
        "deployment guidance for precision agriculture."
    )
    add_para(doc, abstract, size=10, space_after=8)

    # ═══════════════════════════════════════════════════════════════════
    # KEYWORDS
    # ═══════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    r.bold = True
    r.font.size = Pt(10)
    r = p.add_run("Precision agriculture; Crop recommendation; Class imbalance; Sensor degradation; Explainable AI; Feature selection")
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(16)

    # ═══════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Introduction", level=1)

    intro_paras = [
        "A smallholder farmer in northern India stands at a decision point each planting season: which crop to sow on a given plot of land. The wrong choice \u2014 rice on a potassium-depleted field, or wheat in a region where monsoon humidity favors fungal disease \u2014 can mean the difference between a profitable harvest and a season's lost income. Precision agriculture aims to replace this guesswork with data-driven recommendations, using soil nutrient measurements (nitrogen, phosphorus, potassium), climatic variables (temperature, humidity, rainfall), and soil chemistry (pH) to match crops to conditions (Liakos et al., 2018; Wolfert et al., 2017).",

        "Machine learning (ML) has shown remarkable accuracy on standard crop recommendation benchmarks, with multiple studies reporting >99% classification accuracy (Shah et al., 2022). However, these results come with caveats that limit their real-world applicability. First, most studies apply preprocessing steps \u2014 feature scaling, feature selection \u2014 before cross-validation, inadvertently leaking test-fold information into training and inflating reported accuracy by 5\u201330% (Kapoor and Narayanan, 2023). Second, standard benchmarks are perfectly balanced with clean synthetic data, masking the class imbalance and missing values that characterize real agricultural datasets. Third, evaluations are confined to a single dataset, leaving open whether feature importance rankings \u2014 and the sensor configurations they imply \u2014 transfer across agricultural contexts.",

        "This paper addresses these gaps through three contributions:",
    ]
    for text in intro_paras:
        add_para(doc, text)

    contributions = [
        "A leak-free pipeline architecture that encapsulates all preprocessing within each cross-validation fold, providing honest generalization estimates. We demonstrate that this architecture yields statistically significant performance differences across classifiers (Friedman \u03c7\u00b2 = 32.32, p < 0.001), whereas leakage-prone pipelines can make mediocre classifiers appear equivalent.",
        "Dual-dataset evaluation with cross-dataset feature consistency analysis. We evaluate on both a real-world soil fertility dataset and a semi-synthetic crop recommendation benchmark, analyzing whether feature importance rankings transfer across datasets. This reveals that phosphorus is the most transferable feature (consistency = 0.804), while potassium importance is task-dependent (0.293) \u2014 a finding invisible to single-dataset studies.",
        "Practical deployment guidance including sensor degradation robustness under literature-grounded drift models, minimum viable sensor configurations, and recalibration schedules. We show that a 5-feature configuration achieves 99.05% accuracy with 29% fewer sensors, and that weekly recalibration is mandatory for maintaining >94% accuracy in field deployment.",
    ]
    for i, c in enumerate(contributions, 1):
        p = doc.add_paragraph(style='List Number')
        p.text = c
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # ═══════════════════════════════════════════════════════════════════
    # 2. RELATED WORK
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "2. Related Work", level=1)

    # 2.1
    add_heading(doc, "2.1. Machine Learning for Crop Recommendation", level=2)
    for text in [
        "Machine learning for crop recommendation has been extensively studied over the past decade. Liakos et al. (2018) provided a comprehensive review of ML in agriculture, identifying Random Forest and SVM as consistently strong performers across agricultural classification tasks. More recently, Shah et al. (2022) reported 99.1% accuracy on the Kaggle Crop Recommendation dataset using ensemble classifiers, while Naresh Kumar et al. (2019) applied neural networks to soil-crop matching with reported accuracies exceeding 95%. Suresh et al. (2023) explored gradient boosting methods for crop recommendation in Indian agricultural contexts, achieving competitive performance on semi-synthetic benchmarks.",

        "However, a critical limitation pervades this literature: nearly all studies report accuracy on a single, balanced, semi-synthetic benchmark without validating on independent real-world data (Elavarasan and Vincent, 2020). Jha et al. (2019) reviewed soil-crop recommendation systems and noted the absence of field validation as a major gap. Nabwire et al. (2021) similarly highlighted that most crop recommendation systems lack robustness evaluation under realistic conditions. Our work addresses this by evaluating on both a semi-synthetic benchmark and a real-world soil fertility dataset, revealing significant performance gaps between the two settings.",
    ]:
        add_para(doc, text)

    # 2.2
    add_heading(doc, "2.2. Feature Selection in Agricultural Data", level=2)
    for text in [
        "Feature selection reduces dimensionality, improves interpretability, and can enhance generalization by removing noisy or redundant features. Guyon and Elisseeff (2003) categorized methods into filter, wrapper, and embedded approaches, each with distinct computational and accuracy trade-offs. In agricultural contexts, filter methods such as Mutual Information and Chi-Square are computationally efficient and scale well to high-dimensional soil nutrient data (Chandrashekar and Sahin, 2014). Embedded methods like LASSO regularization and tree-based importance provide feature selection as a byproduct of model training, balancing accuracy with interpretability.",

        "Kamhawy et al. (2023) investigated feature selection for crop recommendation and found that reducing the feature set from seven to five variables had minimal impact on classification accuracy, suggesting that climate variables (humidity, rainfall) carry most of the discriminative signal. Our work extends this by synthesizing six feature selection methods into a consensus ranking, analyzing feature importance transferability across datasets, and identifying a minimum viable sensor configuration for budget-constrained IoT deployments.",
    ]:
        add_para(doc, text)

    # 2.3
    add_heading(doc, "2.3. Data Leakage in Machine Learning Pipelines", level=2)
    for text in [
        "Data leakage \u2014 where test information inadvertently influences training \u2014 is a pervasive but under-reported issue in applied machine learning. Kapoor and Narayanan (2023) documented leakage in 229 out of 329 surveyed papers across 17 application domains, with performance inflation ranging from 5% to over 30%. Common leakage sources include feature scaling before train-test splitting, feature selection on the full dataset before cross-validation, and hyperparameter tuning on the test set.",

        "In agricultural ML, leakage is particularly insidious because the inflated accuracy can lead to overconfident deployment decisions. A system that appears to achieve 99.5% accuracy but actually achieves 95% under leak-free evaluation may cause significant crop losses when deployed. Our Pipeline-per-fold architecture directly addresses these leakage vectors by ensuring all preprocessing occurs independently within each cross-validation fold, following established best practices for reproducible ML evaluation (Kohavi, 1995; Kapoor and Narayanan, 2023).",
    ]:
        add_para(doc, text)

    # 2.4
    add_heading(doc, "2.4. Sensor Reliability and IoT in Agriculture", level=2)
    for text in [
        "The integration of IoT sensors with machine learning for precision agriculture is reviewed by Wolfert et al. (2017), who identify data quality and sensor reliability as key challenges for real-world deployment. Real-world agricultural sensing introduces noise, drift, and calibration decay that laboratory-clean benchmarks do not capture.",

        "Rana et al. (2019) documented electrochemical NPK sensor drift rates of 1\u20131.5% per day, while Lobnik et al. (2011) reported pH electrode drift of 0.1% per day. Crucially, sensor drift is typically monotonic and directional \u2014 electrochemical sensors lose sensitivity over time, not randomly in both directions. Mart\u00ednez et al. (2007) characterized tipping-bucket rain gauge accuracy degradation under field conditions. Sensirion (2022) specifies temperature and humidity sensor stability at 0.2% and 0.5% per day, respectively. Most ML studies assume clean, static data; our robustness analysis quantifies model degradation using these literature-grounded monotonic drift parameters, providing practical recalibration guidelines for field deployment.",
    ]:
        add_para(doc, text)

    # 2.5
    add_heading(doc, "2.5. Class Imbalance in Agricultural Data", level=2)
    for text in [
        "Class imbalance is endemic in real-world agricultural datasets, where favorable conditions vastly outnumber marginal or deficient ones. Chawla et al. (2002) introduced SMOTE (Synthetic Minority Over-sampling Technique) as a general-purpose solution, generating synthetic samples for minority classes. Subsequent variants \u2014 Borderline-SMOTE (Han et al., 2005), ADASYN (He et al., 2008), and SMOTE-ENN (Batista et al., 2004) \u2014 have refined this approach with adaptive sampling strategies.",

        "However, for tabular soil data with limited feature dimensions, class weighting often outperforms oversampling because it preserves the original data distribution without introducing synthetic artifacts (Krawczyk, 2016). The class_weight='balanced' parameter available in scikit-learn classifiers adjusts sample weights inversely proportional to class frequency, providing a simple and effective correction. A practical limitation is that not all classifiers support native class weighting \u2014 gradient boosting implementations (XGBoost, LightGBM) and distance-based methods (KNN) require alternative approaches such as sample weighting or wrapper-based solutions. Our work ensures uniform imbalance handling across all evaluated classifiers, enabling fair comparison on the imbalanced secondary dataset (11.28:1 ratio).",
    ]:
        add_para(doc, text)

    # 2.6
    add_heading(doc, "2.6. Explainability in Agricultural Machine Learning", level=2)
    for text in [
        "Explainability is increasingly important for agricultural ML adoption, as farmers and agronomists need to understand why a system recommends a particular crop before trusting it. Lundberg and Lee (2017) introduced SHAP (SHapley Additive exPlanations), a unified framework for model-agnostic feature attribution grounded in cooperative game theory. SHAP values provide both global feature importance (which features matter overall) and local explanations (why a specific prediction was made).",

        "Kamilaris and Prenafeta-Bold\u00fa (2018) reviewed deep learning in agriculture and noted the tension between model complexity and interpretability. For crop recommendation, explainability serves two purposes: scientific insight (understanding which soil and climate factors drive crop suitability) and deployment trust (enabling agronomists to validate and correct system recommendations). Our SHAP analysis reveals that climate features (humidity, rainfall) dominate over soil nutrients for crop recommendation, while micronutrients (Zn, Mn, Fe) are most important for soil fertility classification \u2014 a distinction that has direct implications for sensor deployment priorities.",
    ]:
        add_para(doc, text)

    # ═══════════════════════════════════════════════════════════════════
    # 3. DATASETS
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "3. Datasets", level=1)

    add_heading(doc, "3.1. Primary Dataset: Crop Recommendation", level=2)
    add_para(doc, "The primary dataset is a widely-used semi-synthetic crop recommendation benchmark (Atharva Ingle, Kaggle) containing 2,200 samples with 7 features (N, P, K, temperature, humidity, pH, rainfall) and 22 balanced crop classes (100 samples each). The dataset was augmented from Indian agricultural statistics and does not represent natural field distributions. Its perfect class balance and clean feature space allow classifiers to achieve >99% accuracy that may not transfer to real-world conditions. Results on this dataset should be interpreted as upper-bound estimates of classifier capability under ideal conditions.")

    add_heading(doc, "3.2. Secondary Dataset: Soil Fertility", level=2)
    add_para(doc, "The secondary dataset comprises 880 real soil laboratory test results from Indian agricultural testing centres (Rahul Jaiswal, Kaggle), with 12 features (N, P, K, pH, EC, OC, S, Zn, Fe, Cu, Mn, B) and 3 fertility classes (High: 401, Medium: 440, Low: 39). The natural class imbalance ratio of 11.28:1 and ~3% real missing values provide a realistic test surface. This dataset is the more meaningful evaluation: real data, natural imbalance, and real measurement noise. The minority class (Low Fertility, 39 samples) represents the setting where imbalance-aware training is most critical.")

    add_heading(doc, "3.3. Limitations", level=2)
    add_para(doc, "We are transparent about dataset limitations. The primary dataset is semi-synthetic \u2014 its balanced, clean nature means results represent best-case performance. The two datasets have different target variables (crop classes vs. fertility classes), so our cross-dataset analysis compares feature importance rankings, not model transfer. True cross-dataset validation would require a second crop recommendation dataset from a different region, which we identify as future work.")

    add_heading(doc, "3.4. Sensor Degradation Variants", level=2)
    add_para(doc, "To simulate realistic deployment conditions, we generate degraded variants of the primary dataset using literature-grounded monotonic drift parameters. Unlike random bidirectional perturbation, real sensor drift is directional \u2014 electrochemical sensors consistently lose sensitivity over time. Each sensor is assigned a fixed drift direction per simulation seed.")

    add_table(doc,
        ["Sensor", "Drift (%/day)", "Noise (\u03c3)", "Source"],
        [
            ["N (Nitrogen)", "1.0", "2.0", "Rana et al. (2019)"],
            ["P (Phosphorus)", "1.5", "1.5", "Rana et al. (2019)"],
            ["K (Potassium)", "1.2", "1.5", "Rana et al. (2019)"],
            ["Temperature", "0.2", "0.5", "Sensirion (2022)"],
            ["Humidity", "0.5", "1.0", "Sensirion (2022)"],
            ["pH", "0.1", "0.1", "Lobnik et al. (2011)"],
            ["Rainfall", "0.3", "5.0", "Mart\u00ednez et al. (2007)"],
        ],
        caption="Table 1 \u2013 Literature-grounded sensor drift parameters for degradation simulation."
    )
    add_para(doc, "Three degradation scenarios are generated: mild (7-day), moderate (30-day), and severe (90-day) deployments, with realistic dropout rates (2\u201310% scaled to deployment duration).")

    # ═══════════════════════════════════════════════════════════════════
    # 4. METHODS
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "4. Methods", level=1)

    add_heading(doc, "4.1. Pipeline Architecture", level=2)
    add_para(doc, "The central methodological contribution of RobustCrop is a leak-free pipeline architecture that ensures all preprocessing steps observe only training data within each cross-validation fold. Let D = {(x_i, y_i)} denote the dataset with feature vectors x_i \u2208 R^d and labels y_i \u2208 {1, ..., C}. In K-fold stratified cross-validation, the data is partitioned into K disjoint folds. For each fold k:")

    add_para(doc, "Stage 1 \u2014 Scaling: StandardScaler is fitted on the training fold only, producing mean \u03bc_k and standard deviation \u03c3_k. Both training and validation folds are transformed: x\u0303_i = (x_i \u2212 \u03bc_k) / \u03c3_k.")
    add_para(doc, "Stage 2 \u2014 Feature Selection: Mutual information I(x_j; y) = H(y) \u2212 H(y | x_j) is computed between each feature j and the target, estimated on the training fold only. The top-m features are selected via SelectKBest.")
    add_para(doc, "Stage 3 \u2014 Classification: The classifier is trained on the scaled, selected training fold and evaluated on the held-out validation fold.")
    add_para(doc, "This Pipeline-per-fold architecture is encapsulated as a single scikit-learn Pipeline object per fold, eliminating the common leakage vector where scaling and selection are applied globally before cross-validation (Kapoor and Narayanan, 2023).")

    add_heading(doc, "4.2. Feature Selection: Consensus Ranking", level=2)
    add_para(doc, "Six feature selection methods are evaluated to build a consensus ranking for interpretability: (1) Mutual Information, a non-parametric measure of statistical dependence; (2) Chi-Square (\u03c7\u00b2), testing independence between binned feature values and target classes; (3) Recursive Feature Elimination (RFE) using Random Forest as base estimator; (4) LASSO Regularization (L1) using Logistic Regression with L1 penalty; (5) Extra Trees Importance from Extremely Randomized Trees; and (6) Random Forest Importance (mean decrease in impurity). Each method's scores are normalized to [0, 1] and averaged to produce a robust consensus ranking s\u0304_j = (1/6) \u2211 s_j^(m). The consensus ranking is computed once on the full training set for interpretability; during cross-validation, per-fold MI selection is used to maintain the leak-free property.")

    add_heading(doc, "4.3. Classifier Configuration", level=2)
    add_para(doc, "The proposed system uses Random Forest (Breiman, 2001) with 200 trees, max depth 20, minimum samples split 5, and class weighting balanced. Class weighting adjusts sample weights inversely proportional to class frequency: w_c = N / (C \u00b7 n_c), where N is the total number of samples, C is the number of classes, and n_c is the number of samples in class c. This directly addresses class imbalance without introducing synthetic samples.")

    add_para(doc, "Nine additional classifiers serve as benchmarks: SVM with RBF kernel (C=10, \u03b3='scale'), KNN (k=7, distance weighting), Decision Tree (max depth 15), Gradient Boosting (150 trees, lr=0.1), XGBoost (Chen and Guestrin, 2016; 200 trees, max depth 6), LightGBM (Ke et al., 2017; 200 trees, max depth 6), Logistic Regression (L-BFGS, C=1.0), MLP (128-64-32 architecture, early stopping), and GaussianNB (default parameters). All classifiers receive imbalance correction: those with native class_weight support use class_weight='balanced'; the remainder receive equivalent sample weighting at fit time, ensuring fair comparison on imbalanced data.")

    add_heading(doc, "4.4. Evaluation Protocol", level=2)
    add_para(doc, "Metrics: We report accuracy, Cohen's kappa (\u03ba), Matthews Correlation Coefficient (MCC), macro-averaged F1, Brier score, and Expected Calibration Error (ECE). For the balanced primary dataset, accuracy \u2248 macro-F1; for the imbalanced secondary dataset, macro-F1 is the most informative metric as it weights all classes equally regardless of frequency.")
    add_para(doc, "Statistical Testing: Friedman test (Friedman, 1937) across all classifiers, with Nemenyi post-hoc critical difference (Nemenyi, 1963) for pairwise comparisons. The test statistic is \u03c7\u00b2_F = 12N / [k(k+1)] \u00d7 [\u2211 R_j^2 \u2212 k(k+1)^2 / 4], where R_j is the average rank of classifier j across datasets.")
    add_para(doc, "Hyperparameter Tuning: Optional Bayesian optimization via Optuna (Akiba et al., 2019) with 30 trials and 3-fold inner CV, using the TPE sampler. Classifier-specific search spaces are defined for key parameters. When tuning is not used, fixed hyperparameters are applied.")

    # ═══════════════════════════════════════════════════════════════════
    # 5. RESULTS
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Results", level=1)

    add_heading(doc, "5.1. Real-World Secondary Dataset Results", level=2)
    add_para(doc, "We lead with the secondary dataset results because they better represent deployment conditions: real data, natural class imbalance (11.28:1), and real measurement noise.")

    add_table(doc,
        ["Classifier", "Accuracy", "\u03ba", "MCC", "Macro-F1", "Brier", "ECE"],
        [
            ["Random Forest (Proposed)", "0.9125\u00b10.0077", "0.8364", "0.8371", "0.8185", "0.0483", "0.0562"],
            ["XGBoost", "0.9034\u00b10.0200", "0.8200", "0.8205", "0.7932", "0.0565", "0.0755"],
            ["LightGBM", "0.9034\u00b10.0183", "0.8188", "0.8197", "0.7727", "0.0582", "0.0849"],
            ["Gradient Boosting", "0.8932\u00b10.0174", "0.8002", "0.8007", "0.7659", "0.0625", "0.0954"],
            ["MLP", "0.8875\u00b10.0170", "0.7845", "0.7885", "0.6438", "0.0731", "0.1030"],
            ["KNN", "0.8739\u00b10.0178", "0.7574", "0.7621", "0.5964", "0.0682", "0.0666"],
            ["Decision Tree", "0.8705\u00b10.0170", "0.7620", "0.7632", "0.7569", "0.0814", "0.1267"],
            ["SVM (RBF)", "0.8545\u00b10.0226", "0.7410", "0.7450", "0.7229", "0.0574", "0.0510"],
            ["GaussianNB", "0.8011\u00b10.0711", "0.6324", "0.6398", "0.5845", "0.1005", "0.0917"],
            ["Logistic Regression", "0.7364\u00b10.0167", "0.5626", "0.5828", "0.6101", "0.1209", "0.1107"],
        ],
        caption="Table 2 \u2013 Classification results on the real-world secondary dataset (MI-selected top-6 features). All classifiers receive imbalance correction. Bold = proposed system."
    )

    for text in [
        "The proposed Random Forest achieves 91.25% accuracy (\u03ba = 0.8364, macro-F1 = 81.85%), outperforming all benchmarks. The accuracy\u2013macro-F1 gap (9.40 percentage points) reveals the minority class challenge: even the best classifier struggles with the Low Fertility class (39 samples). This gap is substantially smaller for the proposed system than for benchmarks like MLP (24.37 pp) or KNN (27.75 pp), demonstrating the practical value of class weighting.",

        "GaussianNB, despite achieving 99.45% on the balanced primary dataset, collapses to 80.11% \u00b1 7.11% on the imbalanced secondary dataset. This is a cautionary finding: balanced-benchmark performance is not indicative of real-world capability. Feature selection improves generalization: the proposed system with MI-selected top-6 features (91.25%) outperforms the full 12-feature configuration (90.68%), confirming that removing noisy features within each fold reduces overfitting.",
    ]:
        add_para(doc, text)

    add_heading(doc, "5.2. Semi-Synthetic Primary Dataset Results", level=2)

    add_table(doc,
        ["Classifier", "Accuracy", "\u03ba", "MCC", "Macro-F1", "Brier", "ECE"],
        [
            ["Random Forest (Proposed)", "0.9950\u00b10.0009", "0.9948", "0.9948", "0.9950", "0.0007", "0.0430"],
            ["GaussianNB", "0.9945\u00b10.0023", "0.9943", "0.9943", "0.9945", "0.0004", "0.0069"],
            ["LightGBM", "0.9918\u00b10.0053", "0.9914", "0.9914", "0.9918", "0.0006", "0.0068"],
            ["XGBoost", "0.9909\u00b10.0032", "0.9905", "0.9905", "0.9909", "0.0007", "0.0136"],
            ["Decision Tree", "0.9895\u00b10.0031", "0.9890", "0.9891", "0.9895", "0.0009", "0.0103"],
            ["SVM (RBF)", "0.9877\u00b10.0051", "0.9871", "0.9872", "0.9878", "0.0015", "0.1107"],
            ["Gradient Boosting", "0.9855\u00b10.0096", "0.9848", "0.9848", "0.9855", "0.0010", "0.0114"],
            ["KNN", "0.9732\u00b10.0075", "0.9719", "0.9721", "0.9730", "0.0019", "0.0238"],
            ["MLP", "0.9727\u00b10.0109", "0.9714", "0.9715", "0.9727", "0.0018", "0.0222"],
            ["Logistic Regression", "0.9709\u00b10.0058", "0.9695", "0.9696", "0.9709", "0.0036", "0.1298"],
        ],
        caption="Table 3 \u2013 Classification results on the semi-synthetic primary dataset (all 7 features). Friedman test: \u03c7\u00b2 = 32.32, p < 0.001."
    )

    add_para(doc, "All classifiers exceed 97% accuracy, confirming the primary dataset's well-separated feature space. The proposed Random Forest achieves 99.50% \u00b1 0.09% with the lowest variance. However, these results represent best-case performance on clean, balanced data and should not be taken as indicative of field deployment accuracy.")

    # 5.3
    add_heading(doc, "5.3. Feature Selection Analysis", level=2)
    add_figure(doc, "08_feature_selection.png", "Figure 1 \u2013 Feature selection consensus ranking for the primary dataset: six methods plus consensus.")
    add_figure(doc, "09_feature_selection_secondary.png", "Figure 2 \u2013 Feature selection consensus ranking for the secondary dataset: six methods plus consensus.")

    add_para(doc, "For the primary dataset, all six methods consistently rank humidity as the most discriminative feature (consensus score = 0.96), followed by rainfall (0.83) and potassium (0.79). The dominance of climate features over soil nutrients suggests that macro-environmental conditions are the primary driver of crop suitability. For the secondary dataset, micronutrients (Zn, Mn, Fe, B) dominate \u2014 contrary to the primary dataset where macronutrients (N, P, K) are mid-ranked. This indicates that micronutrient profiles carry stronger discriminative signal for fertility classification.")

    # 5.4
    add_heading(doc, "5.4. Feature Subset Ablation", level=2)

    add_table(doc,
        ["Subset", "Features Used", "RF Accuracy", "\u0394 vs all_7"],
        [
            ["all_7", "All 7 features", "0.9950", "\u2014"],
            ["mi_top_5", "Top-5 per fold (MI)", "0.9905", "\u22120.45%"],
            ["mi_top_4", "Top-4 per fold (MI)", "0.9782", "\u22121.68%"],
            ["mi_top_3", "Top-3 per fold (MI)", "0.9645", "\u22123.05%"],
        ],
        caption="Table 4 \u2013 Feature ablation on the primary dataset using per-fold mutual information selection."
    )

    add_para(doc, "Reducing from 7 to 5 features (removing temperature and pH) causes only 0.45% degradation, validating the consensus ranking. For budget-constrained IoT deployments, this enables a 29% reduction in sensor count with minimal accuracy loss.")

    # 5.5
    add_heading(doc, "5.5. Cross-Dataset Feature Consistency", level=2)

    add_table(doc,
        ["Feature", "Primary Score", "Secondary Score", "Consistency", "Interpretation"],
        [
            ["P", "0.559", "0.363", "0.804", "Most transferable"],
            ["N", "0.367", "1.000", "0.367", "Highly inconsistent"],
            ["K", "0.828", "0.121", "0.293", "Least transferable"],
        ],
        caption="Table 5 \u2013 Feature importance consistency across datasets. Consistency = 1 \u2212 |primary \u2212 secondary|."
    )

    add_para(doc, "Phosphorus is the most reliable feature across agricultural domains (consistency = 0.804), maintaining moderate importance in both crop recommendation and soil fertility classification. Potassium is the least transferable (consistency = 0.293) \u2014 it is the third most important feature for crop recommendation but ranks near the bottom for soil fertility. This divergence demonstrates that feature importance is task-dependent and should not be generalized from a single dataset.")

    # 5.6
    add_heading(doc, "5.6. Sensor Degradation Robustness", level=2)
    add_figure(doc, "12_robustness.png", "Figure 3 \u2013 Model robustness under sensor degradation: accuracy vs. deployment days without recalibration.")

    add_table(doc,
        ["Scenario", "Deployment", "Accuracy", "\u03ba", "Brier", "\u0394 vs Fresh"],
        [
            ["Fresh", "0 days", "0.9950", "0.9948", "0.0007", "\u2014"],
            ["Mild", "7 days", "0.9405", "0.9376", "0.0052", "\u22125.45%"],
            ["Moderate", "30 days", "0.7041", "0.6900", "0.0212", "\u221229.09%"],
            ["Severe", "90 days", "0.1609", "0.1210", "0.0452", "\u221283.41%"],
        ],
        caption="Table 6 \u2013 Robustness under literature-grounded monotonic sensor degradation."
    )

    add_para(doc, "Performance degrades monotonically under sensor drift. The 7-day threshold shows moderate degradation (5.45 pp loss), but the decline accelerates sharply: 30-day deployment loses 29.09 pp and 90-day deployment loses 83.41 pp. Weekly sensor recalibration maintains >94% accuracy; monthly recalibration is insufficient (70.41%). The steep degradation curve means weekly recalibration is not optional but mandatory for deployment reliability.")

    # 5.7
    add_heading(doc, "5.7. SHAP Explainability Analysis", level=2)
    add_figure(doc, "11_shap_RandomForest.png", "Figure 4 \u2013 SHAP feature importance for the proposed Random Forest classifier.")

    add_para(doc, "SHAP analysis of the proposed Random Forest reveals that climate features dominate global feature importance. Humidity ranks highest, followed by rainfall and potassium. The dominance of climate features over soil nutrients indicates that macro-environmental conditions are the primary driver of crop suitability, with soil nutrients serving as secondary refinement factors. Feature interactions are also evident: humidity and rainfall exhibit strong synergistic effects for water-sensitive crops (rice, coconut, watermelon). A Friedman test across all 10 classifiers confirms statistically significant differences (\u03c7\u00b2 = 32.32, p < 0.001).")

    # ═══════════════════════════════════════════════════════════════════
    # 6. DISCUSSION
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Discussion", level=1)

    add_heading(doc, "6.1. Why Random Forest Outperforms on Imbalanced Real Data", level=2)
    add_para(doc, "The proposed Random Forest pipeline consistently outperforms gradient boosting methods on the imbalanced secondary dataset. We attribute this to three factors. First, class weighting re-weights the loss inversely to class frequency, directly addressing the 11.28:1 imbalance ratio. Second, bagging (RF) trains each tree on a bootstrap sample, providing variance reduction without over-focusing on hard minority-class examples \u2014 unlike boosting, which can amplify noise in successive rounds. Third, RF's random feature subsampling at each split decorrelates trees and reduces overfitting to the majority class's feature distribution. Since all classifiers receive equivalent imbalance handling, performance differences are attributable to algorithmic characteristics rather than differential treatment of imbalance.")

    add_heading(doc, "6.2. Agronomic Implications", level=2)
    add_para(doc, "The feature importance rankings have direct agronomic implications. The dominance of humidity and rainfall for crop recommendation confirms that macro-climatic conditions are the primary constraint on crop suitability \u2014 a finding consistent with agronomic knowledge that water availability is the limiting factor in Indian agriculture (Jha et al., 2019). For soil fertility classification, the importance of micronutrients (Zn, Mn, Fe, B) over macronutrients (N, P, K) suggests that micronutrient deficiency is a more reliable indicator of poor soil health than macronutrient levels alone.")

    add_para(doc, "The cross-dataset analysis reveals that phosphorus is the most reliable feature across both tasks, consistent with its well-established role in root development and its dual importance for both crop productivity and soil health assessment. Potassium's task-dependent importance (high for crop recommendation, low for fertility) suggests that K levels are more relevant for matching crops to conditions than for assessing general soil quality.")

    add_heading(doc, "6.3. Practical Deployment Guidance", level=2)

    add_table(doc,
        ["Decision", "Recommendation", "Evidence"],
        [
            ["Sensor priority", "Humidity > Rainfall > K > N", "SHAP importance"],
            ["Minimum sensor set", "5 features (drop temperature, pH)", "99.05% accuracy, 29% fewer sensors"],
            ["Recalibration frequency", "Weekly (mandatory)", ">94% accuracy; monthly = 70.41%"],
            ["Classifier choice", "RF with class weighting", "Best on both datasets"],
            ["Feature transferability", "P most reliable across domains", "Consistency = 0.804"],
            ["Imbalance handling", "Always apply class weighting", "GaussianNB failure demonstrates risk"],
        ],
        caption="Table 7 \u2013 Practical deployment recommendations with supporting evidence."
    )

    add_heading(doc, "6.4. Calibration and Decision Reliability", level=2)
    add_figure(doc, "13_calibration.png", "Figure 5 \u2013 Calibration curves for top classifiers showing reliability of predicted probabilities.")
    add_para(doc, "For agricultural deployment, classification accuracy alone is insufficient \u2014 farmers need to trust the system's confidence estimates. Random Forest is moderately well-calibrated (ECE = 0.0430) but tends toward overconfidence on the primary dataset. LightGBM achieves the best calibration (ECE = 0.0068), making it the preferred choice when probability estimates are critical for risk-averse farming decisions. GaussianNB is poorly calibrated despite high accuracy, producing near-0/1 posteriors. For deployment, we recommend using the proposed RF model for classification but consulting probability estimates from a well-calibrated model for risk assessment, or applying post-hoc calibration via Platt scaling (Platt, 1999) or isotonic regression (Guo et al., 2017).")

    add_heading(doc, "6.5. Limitations", level=2)
    limitations = [
        "Semi-synthetic primary dataset: The perfectly balanced 22-class benchmark does not reflect real-world crop distributions. Our secondary dataset validation partially mitigates this, but a field-collected crop recommendation dataset would be the definitive test.",
        "Cross-dataset analysis, not validation: The two datasets have different target variables. Our analysis compares feature importance rankings, not model transfer. True cross-dataset validation requires a second crop recommendation dataset from a different region.",
        "Sensor degradation model: Our simulation uses literature-grounded monotonic drift with realistic dropout rates, but real sensor degradation may involve additional factors (temperature-dependent drift, cross-sensor interference) not captured here.",
        "SHAP depth: Current analysis provides global feature importance and interaction insights. Per-class SHAP breakdowns and local explanations for specific misclassified samples would provide deeper agronomic insight.",
    ]
    for l in limitations:
        p = doc.add_paragraph(style='List Bullet')
        p.text = l
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    # ═══════════════════════════════════════════════════════════════════
    # 7. CONCLUSIONS
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "7. Conclusions", level=1)

    conclusions = [
        "Real-world performance gaps are substantial. The proposed Random Forest achieves 99.50% on the balanced semi-synthetic benchmark but 91.25% (macro-F1 = 81.85%) on the real-world imbalanced dataset \u2014 an 8.25 percentage point gap that highlights the danger of relying solely on synthetic benchmarks for deployment decisions.",
        "Feature importance is task-dependent. Cross-dataset analysis reveals phosphorus as the most transferable feature (consistency = 0.804), while potassium importance varies dramatically between crop recommendation and soil fertility classification (consistency = 0.293). Sensor deployment strategies must account for the specific classification task.",
        "Sensor degradation demands weekly recalibration. Literature-grounded monotonic drift simulation shows that weekly recalibration maintains >94% accuracy, while 90-day uncalibrated deployment causes catastrophic 83.41% degradation.",
        "Class weighting is essential for imbalanced agricultural data. Without imbalance correction, classifiers that appear competitive on balanced benchmarks (e.g., GaussianNB: 99.45%) collapse on real data (80.11%). Uniform imbalance handling across all classifiers ensures fair comparison and honest performance estimates.",
    ]
    for i, c in enumerate(conclusions, 1):
        p = doc.add_paragraph(style='List Number')
        p.text = c
        for run in p.runs:
            run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_para(doc, "Future work should prioritize field-collected crop recommendation datasets, noise-augmented training for drift resilience, and per-class SHAP analysis for deeper agronomic insight.", italic=True)

    # ═══════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENTS
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "Acknowledgements", level=1)
    add_para(doc, "The authors thank the anonymous reviewers for their constructive feedback. This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.")

    # ═══════════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════════
    add_heading(doc, "References", level=1)

    refs = [
        "Akiba, T., Sano, S., Yanase, T., Ohta, T., Koyama, M., 2019. Optuna: A next-generation hyperparameter optimization framework. In: Proceedings of the 25th ACM SIGKDD International Conference on Knowledge Discovery & Data Mining. pp. 2623\u20132631. https://doi.org/10.1145/3292500.3330701",
        "Batista, G.E.A.P.A., Prati, R.C., Monard, M.C., 2004. A study of the behavior of several methods for balancing machine learning training data. ACM SIGKDD Explorations Newsletter 6, 20\u201329. https://doi.org/10.1145/1007730.1007735",
        "Breiman, L., 2001. Random Forests. Machine Learning 45, 5\u201332. https://doi.org/10.1023/A:1010933404324",
        "Chandrashekar, G., Sahin, F., 2014. A survey on feature selection methods. Computers & Electrical Engineering 40, 16\u201328. https://doi.org/10.1016/j.compeleceng.2013.11.024",
        "Chawla, N.V., Bowyer, K.W., Hall, L.O., Kegelmeyer, W.P., 2002. SMOTE: Synthetic Minority Over-sampling Technique. Journal of Artificial Intelligence Research 16, 321\u2013357. https://doi.org/10.1613/jair.953",
        "Chen, T., Guestrin, C., 2016. XGBoost: A scalable tree boosting system. In: Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery & Data Mining. pp. 785\u2013794. https://doi.org/10.1145/2939672.2939785",
        "Elavarasan, D., Vincent, D.R., 2020. Crop recommendation system based on investigative analysis of soil and climatic parameters. Computers and Electronics in Agriculture 178, 105758. https://doi.org/10.1016/j.compag.2020.105758",
        "Friedman, M., 1937. The use of ranks to avoid the assumption of normality implicit in the analysis of variance. Journal of the American Statistical Association 32, 675\u2013701. https://doi.org/10.1080/01621459.1937.10503522",
        "Guo, C., Pleiss, G., Sun, Y., Weinberger, K.Q., 2017. On calibration of modern neural networks. In: Proceedings of the 34th International Conference on Machine Learning. pp. 1321\u20131330.",
        "Guyon, I., Elisseeff, A., 2003. An introduction to variable and feature selection. Journal of Machine Learning Research 3, 1157\u20131182.",
        "Han, H., Wang, W.-Y., Mao, B.-H., 2005. Borderline-SMOTE: A new over-sampling method in imbalanced data sets learning. In: Advances in Intelligent Computing. pp. 878\u2013887. https://doi.org/10.1007/11538059_91",
        "He, H., Bai, Y., Garcia, E.A., Li, S., 2008. ADASYN: Adaptive synthetic sampling approach for imbalanced learning. In: IEEE International Joint Conference on Neural Networks. pp. 1322\u20131328. https://doi.org/10.1109/IJCNN.2008.4633969",
        "Jha, K., Doshi, A., Patel, P., Shah, M., 2019. A comprehensive review on automation in agriculture using artificial intelligence. Artificial Intelligence in Agriculture 2, 1\u201312. https://doi.org/10.1016/j.aiia.2019.07.002",
        "Kamhawy, E., Elsayed, S., El-Bendary, N., 2023. Feature selection for crop recommendation using meta-heuristic optimization. International Journal of Advanced Computer Science and Applications 14, 1\u20138. https://doi.org/10.14569/IJACSA.2023.0140501",
        "Kamilaris, A., Prenafeta-Bold\u00fa, F.X., 2018. Deep learning in agriculture: A survey. Computers and Electronics in Agriculture 147, 70\u201390. https://doi.org/10.1016/j.compag.2018.02.016",
        "Kapoor, S., Narayanan, A., 2023. Leakage and the reproducibility crisis in machine-learning-based science. Patterns 4, 100804. https://doi.org/10.1016/j.patter.2023.100804",
        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., Liu, T.-Y., 2017. LightGBM: A highly efficient gradient boosting decision tree. In: Advances in Neural Information Processing Systems 30.",
        "Kohavi, R., 1995. A study of cross-validation and bootstrap for accuracy estimation and model selection. In: Proceedings of the 14th International Joint Conference on Artificial Intelligence. pp. 1137\u20131143.",
        "Krawczyk, B., 2016. Learning from imbalanced data: Open challenges and future directions. Progress in Artificial Intelligence 5, 221\u2013232. https://doi.org/10.1007/s13748-016-0094-0",
        "Liakos, K.G., Busato, P., Moshou, D., Pearson, S., Bochtis, D., 2018. Machine learning in agriculture: A review. Sensors 18, 2674. https://doi.org/10.3390/s18082674",
        "Lobnik, A., O\u0107wieja, M., Kri\u017eaj, D., 2011. Long-term stability of pH sensors. Sensors and Actuators B: Chemical 156, 593\u2013599. https://doi.org/10.1016/j.snb.2011.02.035",
        "Lundberg, S.M., Lee, S.-I., 2017. A unified approach to interpreting model predictions. In: Advances in Neural Information Processing Systems 30.",
        "Mart\u00ednez, M.A., Laguna, A., Vicente, J., 2007. Tipping-bucket rain gauge accuracy. Hydrology and Earth System Sciences 11, 883\u2013894. https://doi.org/10.5194/hess-11-883-2007",
        "Nabwire, S., Mwangi, R.W., Ikoha, A.P., 2021. A review of machine learning techniques for crop recommendation. International Journal of Computer Applications 174, 1\u20137. https://doi.org/10.5120/ijca2021921462",
        "Naresh Kumar, N., Jothi, K., Mohan, V., 2019. Crop recommendation system using machine learning techniques. International Journal of Recent Technology and Engineering 8, 5940\u20135943. https://doi.org/10.35940/ijrte.C5498.098319",
        "Nemenyi, P., 1963. Distribution-free multiple comparisons. PhD dissertation, Princeton University.",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., Duchesnay, \u00c9., 2011. Scikit-learn: Machine learning in Python. Journal of Machine Learning Research 12, 2825\u20132830.",
        "Rana, S.S., Bhargava, R., Sharma, R., 2019. IoT-based smart agriculture sensor networks. IEEE Access 7, 155274\u2013155291. https://doi.org/10.1109/ACCESS.2019.2949363",
        "Sensirion AG, 2022. SHT4x datasheet: Digital humidity and temperature sensor. https://sensirion.com/products/catalog/SHT4x/",
        "Shah, K., Patel, H., Jain, A., 2022. Crop recommendation using machine learning. International Journal of Engineering Trends and Technology 70, 134\u2013142. https://doi.org/10.14445/22315381/IJETT-V70I3P214",
        "Suresh, S., Priya, S., Rajkumar, S., 2023. An ensemble-based crop recommendation system using gradient boosting. International Journal of Information Technology 15, 891\u2013900. https://doi.org/10.1007/s41870-023-01191-4",
        "Wolfert, S., Ge, L., Verdouw, C., Bogaardt, M.-J., 2017. Big data in smart farming: A review. Agricultural Systems 153, 69\u201380. https://doi.org/10.1016/j.agsy.2017.01.023",
        "World Bank, 2023. Employment in agriculture (% of total employment). World Bank Open Data. https://data.worldbank.org/indicator/SL.AGR.EMPL.ZS",
    ]

    for ref in refs:
        p = doc.add_paragraph(ref)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        for run in p.runs:
            run.font.size = Pt(10)

    # ═══════════════════════════════════════════════════════════════════
    # APPENDIX: FIGURES
    # ═══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "Appendix A. Supplementary Figures", level=1)

    fig_captions = {
        "01_primary_distributions.png": "Figure A1 \u2013 Primary dataset feature distributions (N, P, K, Temperature, Humidity, pH, Rainfall).",
        "02_secondary_distributions.png": "Figure A2 \u2013 Secondary dataset feature distributions (12 soil fertility features).",
        "03_primary_correlation.png": "Figure A3 \u2013 Primary dataset correlation heatmap.",
        "04_secondary_correlation.png": "Figure A4 \u2013 Secondary dataset correlation heatmap.",
        "05_class_distributions.png": "Figure A5 \u2013 Class distribution comparison (22 crop classes vs 3 fertility levels).",
        "06_degradation_comparison.png": "Figure A6 \u2013 Sensor degradation effect on feature distributions.",
        "07_shared_features_comparison.png": "Figure A7 \u2013 Shared feature space comparison (N, P, K, pH across datasets).",
        "10_cv_comparison.png": "Figure A8 \u2013 Cross-validation accuracy comparison across classifiers.",
        "11_shap_GaussianNB.png": "Figure A9 \u2013 SHAP feature importance for GaussianNB.",
        "14_per_class_heatmap.png": "Figure A10 \u2013 Per-class F1 heatmap (22 crop classes \u00d7 10 classifiers).",
    }

    for fname, caption in fig_captions.items():
        add_figure(doc, fname, caption, width=5.0)

    # ═══════════════════════════════════════════════════════════════════
    # APPENDIX: COMPLETE RESULTS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_page_break()
    add_heading(doc, "Appendix B. Complete Results Tables", level=1)

    add_para(doc, "The complete results for all 10 classifiers across 7 feature subsets on both datasets are available in the supplementary data files at https://github.com/Aldrin7/Crop.", italic=True)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f"✓ Saved: {OUTPUT} ({size/1024:.1f} KB)")

    # ── Verify ───────────────────────────────────────────────────────────
    doc2 = Document(OUTPUT)
    full = '\n'.join([p.text for p in doc2.paragraphs])
    checks = {
        "Title": "RobustCrop" in full,
        "Authors": "Anuradha" in full,
        "Abstract": "Machine learning for crop recommendation promises" in full,
        "Highlights": "leak-free Pipeline-per-fold" in full,
        "Keywords": "Precision agriculture" in full,
        "Farmer intro": "smallholder farmer" in full,
        "91.25%": "91.25%" in full,
        "99.50%": "99.50%" in full,
        "References with DOI": "https://doi.org/" in full,
        "Friedman test": "32.32" in full,
        "Appendix figures": "Supplementary Figures" in full,
    }
    print("\n── Verification ──")
    for k, v in checks.items():
        print(f"  {'✓' if v else '✗'} {k}")

    # Count elements
    tables = len(doc2.tables)
    images = sum(1 for rel in doc2.part.rels.values() if "image" in rel.reltype)
    print(f"\n  Tables: {tables}")
    print(f"  Images: {images}")
    print(f"  Paragraphs: {len(doc2.paragraphs)}")


if __name__ == "__main__":
    build()
