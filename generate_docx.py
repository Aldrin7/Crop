#!/usr/bin/env python3
"""
generate_docx.py — Convert paper_draft_v2.md → RobustCrop_Paper_v2.docx
========================================================================
Publication-quality DOCX with:
  • Proper styles (Heading 1/2, Body Text, Table, Reference, etc.)
  • Tables with borders, bold headers, alternating row shading
  • Unicode math symbols (χ², κ, Σ, ∈, ≥, ≤, etc.)
  • Figure placeholders with captions below; table captions above
  • Numbered references with hanging indent
  • 11pt Times New Roman body, 14pt/12pt bold headings
  • 1.15 line spacing, 0.5″ first-line paragraph indent
  • 1-inch margins all around
  • Running header with short title + page numbers

Usage:
    python generate_docx.py                          # reads paper_draft_v2.md
    python generate_docx.py --input my_paper.md      # custom input
    python generate_docx.py --output MyPaper.docx    # custom output
"""

import re
import sys
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("ERROR: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════
DEFAULT_INPUT = "paper_draft_v2.md"
DEFAULT_OUTPUT = "RobustCrop_Paper_v2.docx"
SHORT_TITLE = "RobustCrop: Leak-Free ML Pipeline for Crop Recommendation"

FONT_BODY = "Times New Roman"
FONT_BODY_SIZE = Pt(11)
FONT_H1_SIZE = Pt(14)
FONT_H2_SIZE = Pt(12)
FONT_H3_SIZE = Pt(11)
LINE_SPACING = 1.15
FIRST_LINE_INDENT = Inches(0.5)
MARGIN = Inches(1)

# Alternating row shading color (light grey)
ALT_ROW_COLOR = "E8E8E8"
HEADER_BG_COLOR = "4472C4"
HEADER_TEXT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)

# Unicode math symbol replacements
MATH_SYMBOLS = {
    r'\chi\^2': 'χ²',
    r'\chi^2': 'χ²',
    r'χ\^2': 'χ²',
    r'chi-squared': 'χ²',
    r'Chi-squared': 'χ²',
    r"Cohen's \\kappa": "Cohen's κ",
    r"Cohen's kappa": "Cohen's κ",
    r'Cohen\'s kappa': "Cohen's κ",
    r'Cohen\'s \\kappa': "Cohen's κ",
    r'\\kappa': 'κ',
    r'\\Sigma': 'Σ',
    r'\\sum': 'Σ',
    r'\\in': '∈',
    r'\\geq': '≥',
    r'\\leq': '≤',
    r'\\pm': '±',
    r'\\times': '×',
    r'\\alpha': 'α',
    r'\\beta': 'β',
    r'\\gamma': 'γ',
    r'\\delta': 'δ',
    r'\\epsilon': 'ε',
    r'\\theta': 'θ',
    r'\\lambda': 'λ',
    r'\\mu': 'μ',
    r'\\sigma': 'σ',
    r'\\tau': 'τ',
    r'\\phi': 'φ',
    r'\\omega': 'ω',
    r'\\pi': 'π',
    r'\\rho': 'ρ',
    r'\\partial': '∂',
    r'\\nabla': '∇',
    r'\\infty': '∞',
    r'\\approx': '≈',
    r'\\neq': '≠',
    r'\\propto': '∝',
    r'\\rightarrow': '→',
    r'\\leftarrow': '←',
    r'\\Rightarrow': '⇒',
    r'\\forall': '∀',
    r'\\exists': '∃',
}


# ═══════════════════════════════════════════════════════════════════════════════
# STYLE SETUP
# ═══════════════════════════════════════════════════════════════════════════════
def setup_styles(doc):
    """Create/override document styles for publication quality."""
    styles = doc.styles

    # --- Body Text ---
    style_body = styles['Normal']
    _set_font(style_body, FONT_BODY, FONT_BODY_SIZE)
    pf = style_body.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.first_line_indent = FIRST_LINE_INDENT

    # --- Heading 1 ---
    h1 = styles['Heading 1']
    _set_font(h1, FONT_BODY, FONT_H1_SIZE, bold=True, color=RGBColor(0x1A, 0x1A, 0x5C))
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Inches(0)
    h1.paragraph_format.keep_with_next = True

    # --- Heading 2 ---
    h2 = styles['Heading 2']
    _set_font(h2, FONT_BODY, FONT_H2_SIZE, bold=True, color=RGBColor(0x2C, 0x2C, 0x6C))
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.first_line_indent = Inches(0)
    h2.paragraph_format.keep_with_next = True

    # --- Heading 3 ---
    h3 = styles['Heading 3']
    _set_font(h3, FONT_BODY, FONT_H3_SIZE, bold=True, color=RGBColor(0x3D, 0x3D, 0x7C))
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.first_line_indent = Inches(0)
    h3.paragraph_format.keep_with_next = True

    # --- Table Caption (above tables) ---
    if 'Table Caption' not in [s.name for s in styles]:
        tc = styles.add_style('Table Caption', 1)  # 1 = paragraph style
    else:
        tc = styles['Table Caption']
    _set_font(tc, FONT_BODY, Pt(10), bold=True, italic=True)
    tc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tc.paragraph_format.space_before = Pt(12)
    tc.paragraph_format.space_after = Pt(4)
    tc.paragraph_format.first_line_indent = Inches(0)

    # --- Figure Caption (below figures) ---
    if 'Figure Caption' not in [s.name for s in styles]:
        fc = styles.add_style('Figure Caption', 1)
    else:
        fc = styles['Figure Caption']
    _set_font(fc, FONT_BODY, Pt(10), italic=True)
    fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fc.paragraph_format.space_before = Pt(4)
    fc.paragraph_format.space_after = Pt(12)
    fc.paragraph_format.first_line_indent = Inches(0)

    # --- Reference style ---
    if 'Reference' not in [s.name for s in styles]:
        ref = styles.add_style('Reference', 1)
    else:
        ref = styles['Reference']
    _set_font(ref, FONT_BODY, Pt(10))
    ref.paragraph_format.first_line_indent = Inches(-0.3)
    ref.paragraph_format.left_indent = Inches(0.3)
    ref.paragraph_format.space_after = Pt(3)
    ref.paragraph_format.line_spacing = LINE_SPACING

    # --- Code / Equation block ---
    if 'Equation' not in [s.name for s in styles]:
        eq = styles.add_style('Equation', 1)
    else:
        eq = styles['Equation']
    _set_font(eq, "Cambria Math", Pt(11))
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(6)
    eq.paragraph_format.space_after = Pt(6)
    eq.paragraph_format.first_line_indent = Inches(0)

    # --- Abstract style ---
    if 'Abstract' not in [s.name for s in styles]:
        ab = styles.add_style('Abstract', 1)
    else:
        ab = styles['Abstract']
    _set_font(ab, FONT_BODY, Pt(10), italic=True)
    ab.paragraph_format.left_indent = Inches(0.5)
    ab.paragraph_format.right_indent = Inches(0.5)
    ab.paragraph_format.first_line_indent = Inches(0)
    ab.paragraph_format.space_after = Pt(6)


def _set_font(style, name, size, bold=False, italic=False, color=None):
    """Helper to set font properties on a style."""
    font = style.font
    font.name = name
    font.size = size
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color
    # Set East Asian font
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), name)


# ═══════════════════════════════════════════════════════════════════════════════
# HEADER / FOOTER
# ═══════════════════════════════════════════════════════════════════════════════
def setup_header_footer(doc):
    """Add running header with short title and page numbers in footer."""
    for section in doc.sections:
        # Margins
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

        # Header
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = SHORT_TITLE
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.runs[0] if hp.runs else hp.add_run()
        run.font.size = Pt(9)
        run.font.name = FONT_BODY
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        # Add bottom border to header paragraph
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="808080"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Footer with page numbers
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clear existing
        for run in fp.runs:
            run.text = ""
        # Add "Page X of Y"
        run1 = fp.add_run("Page ")
        run1.font.size = Pt(9)
        run1.font.name = FONT_BODY
        run1.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # PAGE field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_field = fp.add_run()
        run_field._r.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_instr = fp.add_run()
        run_instr._r.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end = fp.add_run()
        run_end._r.append(fldChar2)

        run2 = fp.add_run(" of ")
        run2.font.size = Pt(9)
        run2.font.name = FONT_BODY
        run2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        # NUMPAGES field
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_field2 = fp.add_run()
        run_field2._r.append(fldChar3)
        instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_instr2 = fp.add_run()
        run_instr2._r.append(instrText2)
        fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_end2 = fp.add_run()
        run_end2._r.append(fldChar4)


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE CREATION
# ═══════════════════════════════════════════════════════════════════════════════
def add_table(doc, headers, rows, caption=None):
    """Add a formatted table with borders, bold headers, alternating rows."""
    if caption:
        p = doc.add_paragraph(style='Table Caption')
        p.add_run(caption)

    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Apply borders to entire table
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    hdr_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = FONT_BODY
        run.font.color.rgb = HEADER_TEXT_COLOR
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{HEADER_BG_COLOR}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows with alternating shading
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            run.font.name = FONT_BODY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating row shading
            if row_idx % 2 == 1:
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="{ALT_ROW_COLOR}" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # spacing after table
    return table


# ═══════════════════════════════════════════════════════════════════════════════
# MATH SYMBOL PROCESSING
# ═══════════════════════════════════════════════════════════════════════════════
def replace_math_symbols(text):
    """Replace LaTeX-style math with Unicode symbols."""
    for pattern, replacement in MATH_SYMBOLS.items():
        text = text.replace(pattern, replacement)
    # Inline math: $...$ → just the content with symbol replacements
    text = re.sub(r'\$([^$]+)\$', lambda m: _process_inline_math(m.group(1)), text)
    return text


def _process_inline_math(math_text):
    """Process inline math expression, replacing LaTeX commands with Unicode."""
    result = math_text
    for pattern, replacement in MATH_SYMBOLS.items():
        result = result.replace(pattern, replacement)
    # Clean remaining LaTeX
    result = re.sub(r'\\[a-zA-Z]+', '', result)
    result = result.replace('{', '').replace('}', '')
    return result.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# MARKDOWN PARSER
# ═══════════════════════════════════════════════════════════════════════════════
def parse_markdown(md_text):
    """Parse markdown into a list of block elements.

    Returns list of dicts:
      {'type': 'h1'|'h2'|'h3'|'p'|'table'|'figure'|'equation'|'ref_list'|'code'|'abstract',
       'content': ...}
    """
    lines = md_text.split('\n')
    blocks = []
    i = 0
    current_para = []

    def flush_para():
        nonlocal current_para
        if current_para:
            text = '\n'.join(current_para).strip()
            if text:
                # Check for special blocks
                if text.lower().startswith('abstract') and len(text) > 9:
                    blocks.append({'type': 'abstract', 'content': text})
                else:
                    blocks.append({'type': 'p', 'content': text})
            current_para = []

    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith('### '):
            flush_para()
            blocks.append({'type': 'h3', 'content': line[4:].strip()})
            i += 1
            continue
        if line.startswith('## '):
            flush_para()
            blocks.append({'type': 'h2', 'content': line[3:].strip()})
            i += 1
            continue
        if line.startswith('# ') and not line.startswith('## '):
            flush_para()
            blocks.append({'type': 'h1', 'content': line[2:].strip()})
            i += 1
            continue

        # Code block (equation or code)
        if line.strip().startswith('```'):
            flush_para()
            lang = line.strip().lstrip('`').strip()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = '\n'.join(code_lines)
            if lang in ('math', 'latex', 'equation'):
                blocks.append({'type': 'equation', 'content': code_text})
            else:
                blocks.append({'type': 'code', 'content': code_text})
            continue

        # Table (line starts with |)
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            flush_para()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'table', 'content': table_lines})
            continue

        # Figure reference: ![caption](path) or Figure N:
        fig_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
        if fig_match:
            flush_para()
            blocks.append({
                'type': 'figure',
                'caption': fig_match.group(1),
                'path': fig_match.group(2),
            })
            i += 1
            continue

        # Figure caption line: "Figure N:" or "Fig. N:"
        fig_cap_match = re.match(r'^(Fig(?:ure)?\.?\s*\d+[.:]\s*.+)', line.strip(), re.IGNORECASE)
        if fig_cap_match:
            flush_para()
            blocks.append({'type': 'figure_caption', 'content': fig_cap_match.group(1)})
            i += 1
            continue

        # Table caption line: "Table N:" or "Tab. N:"
        tab_cap_match = re.match(r'^(Tab(?:le)?\.?\s*\d+[.:]\s*.+)', line.strip(), re.IGNORECASE)
        if tab_cap_match:
            flush_para()
            blocks.append({'type': 'table_caption', 'content': tab_cap_match.group(1)})
            i += 1
            continue

        # Reference section detection
        if re.match(r'^#+\s*References?\s*$', line.strip(), re.IGNORECASE):
            flush_para()
            blocks.append({'type': 'h1', 'content': 'References'})
            i += 1
            # Collect reference lines
            ref_lines = []
            while i < len(lines):
                ref_line = lines[i].strip()
                if not ref_line:
                    i += 1
                    continue
                # Numbered reference: [1] ... or 1. ...
                if re.match(r'^\[?\d+\]?[\.\)]\s', ref_line) or re.match(r'^\[\d+\]', ref_line):
                    ref_lines.append(ref_line)
                elif ref_line.startswith('#'):
                    break  # next section
                else:
                    ref_lines.append(ref_line)
                i += 1
            if ref_lines:
                blocks.append({'type': 'ref_list', 'content': ref_lines})
            continue

        # Empty line → flush paragraph
        if not line.strip():
            flush_para()
            i += 1
            continue

        # Regular text → accumulate into paragraph
        current_para.append(line)
        i += 1

    flush_para()
    return blocks


# ═══════════════════════════════════════════════════════════════════════════════
# RICH TEXT RENDERING (bold, italic, inline code, links)
# ═══════════════════════════════════════════════════════════════════════════════
def add_rich_paragraph(doc, text, style='Normal', alignment=None):
    """Add a paragraph with inline formatting (**bold**, *italic*, `code`, [link](url))."""
    text = replace_math_symbols(text)
    p = doc.add_paragraph(style=style)
    if alignment:
        p.alignment = alignment

    # Pattern for inline formatting
    # Order: bold+italic, bold, italic, inline code, links, plain text
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'       # ***bold italic***
        r'|(\*\*(.+?)\*\*)'           # **bold**
        r'|(\*([^*]+?)\*)'            # *italic*
        r'|(`([^`]+?)`)'              # `code`
        r'|\[([^\]]+)\]\(([^)]+)\)'   # [text](url)
        r'|([^*_`\[]+)'               # plain text
    )

    for m in pattern.finditer(text):
        if m.group(2):  # bold italic
            run = p.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(4):  # bold
            run = p.add_run(m.group(4))
            run.bold = True
        elif m.group(6):  # italic
            run = p.add_run(m.group(6))
            run.italic = True
        elif m.group(8):  # code
            run = p.add_run(m.group(8))
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            # Light grey background for inline code
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
            run._r.get_or_add_rPr().append(shading)
        elif m.group(9):  # link
            run = p.add_run(m.group(9))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.underline = True
        elif m.group(11):  # plain
            run = p.add_run(m.group(11))

        # Set default font for all runs
        if hasattr(run, 'font'):
            run.font.name = FONT_BODY
            run.font.size = FONT_BODY_SIZE

    return p


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE RENDERING
# ═══════════════════════════════════════════════════════════════════════════════
def render_table(table_lines):
    """Parse markdown table lines into headers and rows."""
    if len(table_lines) < 2:
        return None, None

    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    headers = parse_row(table_lines[0])
    # Skip separator line (line with dashes)
    rows = []
    for line in table_lines[2:]:
        if re.match(r'^[\s|:\-]+$', line):
            continue
        row = parse_row(line)
        # Pad or trim to match header count
        while len(row) < len(headers):
            row.append('')
        rows.append(row[:len(headers)])

    return headers, rows


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ═══════════════════════════════════════════════════════════════════════════════
def build_document(md_text, output_path):
    """Convert parsed markdown into a DOCX document."""
    doc = Document()

    # ── Setup styles and page layout ──
    setup_styles(doc)
    setup_header_footer(doc)

    # ── Parse markdown ──
    blocks = parse_markdown(md_text)

    # Track section numbering
    section_counter = [0, 0, 0]  # h1, h2, h3
    in_references = False
    pending_table_caption = None

    for block in blocks:
        btype = block['type']
        content = block.get('content', '')

        # ── Headings ──
        if btype == 'h1':
            if content.lower() == 'references':
                in_references = True
                p = doc.add_paragraph(style='Heading 1')
                p.add_run('References')
            else:
                in_references = False
                section_counter[0] += 1
                section_counter[1] = 0
                section_counter[2] = 0
                num = f"{section_counter[0]}"
                p = doc.add_paragraph(style='Heading 1')
                p.add_run(f"{num}. {replace_math_symbols(content)}")
            continue

        if btype == 'h2':
            section_counter[1] += 1
            section_counter[2] = 0
            num = f"{section_counter[0]}.{section_counter[1]}"
            p = doc.add_paragraph(style='Heading 2')
            p.add_run(f"{num} {replace_math_symbols(content)}")
            continue

        if btype == 'h3':
            section_counter[2] += 1
            num = f"{section_counter[0]}.{section_counter[1]}.{section_counter[2]}"
            p = doc.add_paragraph(style='Heading 3')
            p.add_run(f"{num} {replace_math_symbols(content)}")
            continue

        # ── Abstract ──
        if btype == 'abstract':
            # Remove "Abstract:" or "Abstract" prefix
            text = re.sub(r'^Abstract\s*:?\s*', '', content, flags=re.IGNORECASE)
            p = doc.add_paragraph(style='Abstract')
            run = p.add_run(replace_math_symbols(text))
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            run.italic = True
            continue

        # ── Table caption (stored for next table) ──
        if btype == 'table_caption':
            pending_table_caption = replace_math_symbols(content)
            continue

        # ── Table ──
        if btype == 'table':
            headers, rows = render_table(content)
            if headers and rows:
                caption = pending_table_caption
                pending_table_caption = None
                add_table(doc, headers, rows, caption=caption)
            continue

        # ── Figure ──
        if btype == 'figure':
            caption = block.get('caption', '')
            # Add placeholder paragraph
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[Figure: {caption}]')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
            run.italic = True
            # Add figure caption below
            if caption:
                fc = doc.add_paragraph(style='Figure Caption')
                fc.add_run(replace_math_symbols(caption))
            continue

        if btype == 'figure_caption':
            p = doc.add_paragraph(style='Figure Caption')
            p.add_run(replace_math_symbols(content))
            continue

        # ── Equation ──
        if btype == 'equation':
            eq_text = replace_math_symbols(content)
            p = doc.add_paragraph(style='Equation')
            run = p.add_run(eq_text)
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            continue

        # ── Code block ──
        if btype == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # Light grey background for code blocks
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
            run._r.get_or_add_rPr().append(shading)
            continue

        # ── Reference list ──
        if btype == 'ref_list':
            for ref_text in content:
                ref_text = replace_math_symbols(ref_text)
                # Extract reference number
                match = re.match(r'^\[?(\d+)\]?\s*[\.\):]?\s*(.+)', ref_text)
                if match:
                    num = match.group(1)
                    ref_body = match.group(2)
                    p = doc.add_paragraph(style='Reference')
                    run_num = p.add_run(f'[{num}] ')
                    run_num.bold = True
                    run_num.font.name = FONT_BODY
                    run_num.font.size = Pt(10)
                    run_body = p.add_run(ref_body)
                    run_body.font.name = FONT_BODY
                    run_body.font.size = Pt(10)
                else:
                    p = doc.add_paragraph(style='Reference')
                    run = p.add_run(ref_text)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
            continue

        # ── Regular paragraph ──
        if btype == 'p':
            if in_references:
                # In references section, use reference style
                ref_text = replace_math_symbols(content)
                match = re.match(r'^\[?(\d+)\]?\s*[\.\):]?\s*(.+)', ref_text)
                if match:
                    num = match.group(1)
                    ref_body = match.group(2)
                    p = doc.add_paragraph(style='Reference')
                    run_num = p.add_run(f'[{num}] ')
                    run_num.bold = True
                    run_num.font.name = FONT_BODY
                    run_num.font.size = Pt(10)
                    run_body = p.add_run(ref_body)
                    run_body.font.name = FONT_BODY
                    run_body.font.size = Pt(10)
                else:
                    p = doc.add_paragraph(style='Reference')
                    run = p.add_run(ref_text)
                    run.font.name = FONT_BODY
                    run.font.size = Pt(10)
            else:
                add_rich_paragraph(doc, content)
            continue

    # ── Save ──
    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(
        description='Convert paper_draft_v2.md to publication-quality DOCX')
    parser.add_argument('--input', '-i', default=DEFAULT_INPUT,
                        help=f'Input markdown file (default: {DEFAULT_INPUT})')
    parser.add_argument('--output', '-o', default=DEFAULT_OUTPUT,
                        help=f'Output DOCX file (default: {DEFAULT_OUTPUT})')
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    if not input_path.exists():
        print(f"ERROR: Input file '{input_path}' not found.")
        print(f"  Expected location: {input_path.resolve()}")
        print("  Create paper_draft_v2.md first, then run this script.")
        sys.exit(1)

    print(f"Reading: {input_path}")
    md_text = input_path.read_text(encoding='utf-8')
    print(f"  {len(md_text)} characters, {md_text.count(chr(10))} lines")

    print(f"Generating: {output_path}")
    build_document(md_text, str(output_path))

    size_kb = output_path.stat().st_size / 1024
    print(f"  ✓ Saved: {output_path} ({size_kb:.1f} KB)")
    print("Done.")


if __name__ == '__main__':
    main()
